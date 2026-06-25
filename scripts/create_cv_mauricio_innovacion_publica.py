from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"D:\development\javascript\observatorio\CV_Mauricio_Ojeda_Innovacion_Publica_2026.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)
LIGHT_FILL = "F2F4F7"
BORDER = "D9E2EF"


def set_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[index]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_para(paragraph, before=0, after=6, line=1.10, align=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    style_para(p, before=16 if level == 1 else 10, after=6)
    r = p.add_run(text)
    if level == 1:
        set_font(r, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_font(r, size=13, color=BLUE, bold=True)
    else:
        set_font(r, size=12, color=DARK_BLUE, bold=True)
    return p


def add_body(doc, text, after=6, bold_lead=None):
    p = doc.add_paragraph()
    style_para(p, after=after)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest)
    else:
        set_font(p.add_run(text))
    return p


def add_bullet(doc, text, after=4):
    p = doc.add_paragraph(style="List Bullet")
    style_para(p, after=after, line=1.167)
    for run in p.runs:
        set_font(run)
    set_font(p.add_run(text))
    return p


def add_role(doc, org, role, dates, bullets):
    p = doc.add_paragraph()
    style_para(p, before=8, after=2)
    set_font(p.add_run(org), size=12, color=DARK_BLUE, bold=True)
    p2 = doc.add_paragraph()
    style_para(p2, after=4)
    set_font(p2.add_run(role), bold=True)
    set_font(p2.add_run(f" | {dates}"), color=MUTED)
    for item in bullets:
        add_bullet(doc, item)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2500, 6860])
    set_table_borders(table)
    for i, (label, value) in enumerate(rows):
        cells = table.rows[i].cells
        set_cell_shading(cells[0], LIGHT_FILL)
        for paragraph in cells[0].paragraphs:
            style_para(paragraph, after=0)
            set_font(paragraph.add_run(label), bold=True, color=DARK_BLUE)
        for paragraph in cells[1].paragraphs:
            style_para(paragraph, after=0)
            set_font(paragraph.add_run(value))
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)

for style_name in ("List Bullet", "List Number"):
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(11)

header = section.header.paragraphs[0]
style_para(header, after=0)
set_font(header.add_run("CV orientado a innovación pública digital"), size=9, color=MUTED)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
style_para(footer, after=0)
set_font(footer.add_run("Mauricio Ojeda Pepinosa"), size=9, color=MUTED)

title = doc.add_paragraph()
style_para(title, before=8, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(title.add_run("MAURICIO OJEDA PEPINOSA"), size=22, color=INK, bold=True)

subtitle = doc.add_paragraph()
style_para(subtitle, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(
    subtitle.add_run("Ingeniero Industrial | Ciencia de datos | Sistemas de información e innovación pública"),
    size=12,
    color=MUTED,
    bold=True,
)

contact = doc.add_paragraph()
style_para(contact, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(contact.add_run("Bogota D.C., Colombia | jmojedap@gmail.com"), size=10.5, color=MUTED)

add_key_value_table(
    doc,
    [
        ("Rol objetivo", "Investigador encargado del procesamiento, análisis y visualización de datos."),
        (
            "Convocatoria",
            "Iniciativas de innovación pública con enfoque digital centrado en las personas.",
        ),
        (
            "Aporte diferencial",
            "Integración de investigación aplicada, sistemas de información, modelamiento de datos, analítica y visualización para evidenciar resultados e impacto público.",
        ),
    ],
)

add_heading(doc, "Perfil profesional", 1)
add_body(
    doc,
    "Ingeniero Industrial con experiencia en captura, procesamiento, análisis y visualización de datos, así como en el diseño y desarrollo de sistemas de información para organizaciones públicas, sociales y privadas. Ha liderado y acompañado procesos de estructuración de modelos de datos, herramientas de recolección, interfaces interactivas, tableros y soluciones digitales orientadas a la gestión del conocimiento, la cultura ciudadana, la educación y el seguimiento de servicios públicos.",
)
add_body(
    doc,
    "Su perfil combina investigación aplicada, analítica estadística, desarrollo de software y comprensión de problemas institucionales. Puede aportar evidencia técnica para documentar resultados, medir impacto, mejorar procesos y fortalecer iniciativas digitales centradas en las personas.",
)

add_heading(doc, "Competencias clave para la postulación", 1)
for item in [
    "Procesamiento, depuración, integración y análisis de datos institucionales.",
    "Diseño de metodologías, instrumentos e interfaces de captura de información.",
    "Construcción de tableros, visualizaciones y herramientas de consulta dinámica de resultados.",
    "Análisis estadístico, segmentación, pronóstico de series de tiempo y modelamiento de variables.",
    "Levantamiento, definición y documentación de requerimientos funcionales para sistemas de información.",
    "Desarrollo de soluciones web y bases de datos para proyectos de impacto público y social.",
    "Traducción de hallazgos técnicos en insumos útiles para toma de decisiones, evaluación y comunicación institucional.",
]:
    add_bullet(doc, item)

add_heading(doc, "Experiencia relevante", 1)
add_role(
    doc,
    "Secretaría de Cultura, Recreación y Deporte - Bogotá D.C.",
    "Contratista, Observatorio de Cultura: Líder de sistemas de información y visualización de datos",
    "Febrero de 2022 - Actualidad",
    [
        "Orienta y acompaña el diseño y desarrollo del módulo de analítica de datos y flujo de información del Sistema de Información Misional.",
        "Define y documenta requerimientos y funcionalidades relacionadas con cultura ciudadana y gestión del conocimiento cultural.",
        "Diseña e implementa metodologías, herramientas e interfaces de recolección, análisis y visualización de datos.",
        "Aporta a la transformación de información institucional en evidencia para seguimiento, consulta y toma de decisiones.",
    ],
)

add_role(
    doc,
    "Pacarina Media Lab",
    "Socio, desarrollador y analista",
    "Septiembre de 2009 - Actualidad",
    [
        "Desarrolla soluciones digitales para captura, procesamiento, consulta y visualización interactiva de información.",
        "Proyecto Todos En Sintonía con la Justicia, para Cordupaz, USAID y Ministerio de Justicia: interfaz interactiva de captura de datos y herramienta de consulta dinámica de resultados.",
        "Proyecto El Color de Tu Verdad, para Pacifista y Comisión de la Verdad: interfaz interactiva de captura de datos y herramienta de consulta dinámica de resultados.",
        "Proyecto para el Ministerio de Educación Nacional: diseño, desarrollo e implementación de plataforma e-learning para los Puntos Vive Digital de la Región Caribe.",
        "Proyecto Metrofútbol: análisis estadístico de variables cuantitativas, cualitativas y espaciales del fútbol profesional colombiano.",
        "Revista Minutos de Amor: diseño y desarrollo de algoritmo optimizado para cálculo mixto de unidades de empaque en despacho de pedidos.",
    ],
)

add_role(
    doc,
    "Empresas Públicas de Medellín E.S.P.",
    "Analista de consumo, Dirección Aguas - Área de Normalización y Soporte",
    "Mayo de 2008 - Julio de 2009",
    [
        "Realizó análisis estadístico y pronóstico de series de tiempo para variables de consumo de servicios de acueducto y alcantarillado en Medellín y el Área Metropolitana del Valle de Aburrá.",
        "Desarrolló e implementó un modelo matemático para el cálculo y seguimiento de variables de consumo.",
        "Efectuó segmentación, análisis, seguimiento y proyección de demanda por estrato, clase de uso, circuitos hidráulicos y otros criterios operativos.",
    ],
)

add_role(
    doc,
    "EPM Bogotá Aguas S.A. E.S.P.",
    "Analista de la demanda, Dirección Comercial - Mercadeo y Nuevos Negocios",
    "Enero de 2007 - Abril de 2008",
    [
        "Analizó y realizó seguimiento del consumo de agua en Bogotá.",
        "Desarrolló herramientas informáticas para seguimiento de demanda por estrato, clase de uso, ciclo de facturación y porción geográfica.",
        "Participó en el diseño de procesos de optimización de ingresos.",
    ],
)

add_role(
    doc,
    "EPM Bogotá Aguas S.A. E.S.P.",
    "Desarrollador de aplicaciones MS Access para interventoría y gestión comercial",
    "Experiencia previa",
    [
        "Desarrolló aplicaciones para procesos de interventoría, seguimiento del cumplimiento y evaluación de la gestión comercial del contratista de la EAAB.",
    ],
)

add_heading(doc, "Herramientas y capacidades técnicas", 1)
add_key_value_table(
    doc,
    [
        ("Lenguajes", "JavaScript, PHP, SQL, HTML, CSS, Python, Dart."),
        ("Frameworks y librerías", "CodeIgniter, Vue.js, Bootstrap, Flutter, Jupyter, Pandas."),
        ("Bases de datos", "MySQL, MS Access."),
        ("Ofimática y análisis", "Excel avanzado, análisis estadístico, modelos de datos, visualización de información."),
    ],
)

add_heading(doc, "Educación", 1)
add_body(doc, "Universidad de La Salle | Especialización en Gerencia de Mercadeo | 2009-2010", bold_lead="Universidad de La Salle")
add_body(
    doc,
    "Universidad Nacional de Colombia | Ingeniero Industrial | 1999-2004 | Línea de profundización: Sistemas de Información para las organizaciones.",
    bold_lead="Universidad Nacional de Colombia",
)

add_heading(doc, "Reconocimientos académicos", 1)
add_body(
    doc,
    "Bachilleres por Colombia - 5to puesto | Ganador Beca Ecopetrol, Programa Bachilleres por Colombia - Mario Galán Gómez | 2000.",
)
add_body(
    doc,
    "Beca Andrés Bello, ICETEX | Categoría Departamental, 1er puntaje ICFES Nariño | 1999.",
)

add_heading(doc, "Idiomas", 1)
add_body(doc, "Inglés: lectura 80%; conversación 25%.")

add_heading(doc, "Referencias", 1)
add_body(doc, "Referencias laborales disponibles a solicitud.")

doc.save(OUTPUT)
print(OUTPUT)
